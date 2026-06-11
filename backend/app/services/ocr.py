import json
import os
import re
import shutil
import struct
import subprocess
import tempfile
import time
import unicodedata
from decimal import Decimal, InvalidOperation
from pathlib import Path

import httpx
from docx import Document
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.models.entities import OcrRecognitionLog
from app.services.system_config import LLMRuntimeConfig, get_llm_runtime_config


_CN_DIGITS = {
    "零": 0,
    "〇": 0,
    "一": 1,
    "壹": 1,
    "二": 2,
    "贰": 2,
    "两": 2,
    "三": 3,
    "叁": 3,
    "四": 4,
    "肆": 4,
    "五": 5,
    "伍": 5,
    "六": 6,
    "陆": 6,
    "七": 7,
    "柒": 7,
    "八": 8,
    "捌": 8,
    "九": 9,
    "玖": 9,
}
_CN_UNITS = {"十": 10, "拾": 10, "百": 100, "佰": 100, "千": 1000, "仟": 1000}


def _normalize_text(text: str) -> str:
    text = unicodedata.normalize("NFKC", text or "")
    return text.replace("\u3000", " ").replace("\r\n", "\n").replace("\r", "\n")


def _clean_value(value: str | None, limit: int = 120) -> str:
    if not value:
        return ""
    value = unicodedata.normalize("NFKC", value)
    value = re.split(r"[\n\r;；]", value, 1)[0]
    value = re.split(r"[（(]以下简称", value, 1)[0]
    value = re.sub(r"\s+", " ", value)
    return value.strip(" \t:：,，.。)）(（[]【】\"'").strip()[:limit]


def _decimal_text(value: str | None) -> str:
    if not value:
        return ""
    try:
        normalized = value.replace(",", "").strip()
        if normalized.count(".") > 1:
            parts = normalized.split(".")
            normalized = "".join(parts[:-1]) + "." + parts[-1]
        return f"{Decimal(normalized).quantize(Decimal('0.01'))}"
    except (InvalidOperation, AttributeError):
        return ""


def _to_decimal(value: str | None) -> Decimal | None:
    amount = _decimal_text(value)
    if not amount:
        return None
    try:
        return Decimal(amount)
    except InvalidOperation:
        return None


def _money_values(text: str) -> list[str]:
    values: list[str] = []
    for match in re.finditer(r"([¥￥]\s*)?([0-9][0-9,\.]*\d)", _normalize_text(text)):
        raw = match.group(2)
        if re.match(r"20\d{2}[./-]\d{1,2}[./-]\d{1,2}$", raw):
            continue
        if not match.group(1) and "." not in raw and "," not in raw:
            continue
        amount = _decimal_text(raw)
        if not amount:
            continue
        value = Decimal(amount)
        # 票据代码、税号、电话、日期等没有小数或金额过小，排除掉避免污染候选。
        if value >= Decimal("100"):
            values.append(amount)
    return values


def _extract_amount_by_labels(text: str, labels: list[str], window: int = 160, prefer: str = "first") -> str:
    text = _normalize_text(text)
    for label in labels:
        for match in re.finditer(re.escape(label), text, re.I):
            snippet = text[match.end(): match.end() + window]
            values = _money_values(snippet)
            if values:
                if prefer == "max":
                    return max(values, key=Decimal)
                if prefer == "last":
                    return values[-1]
                return values[0]
    return ""


def _line_values(text: str) -> list[str]:
    return [line.strip() for line in re.split(r"[\n\r]+", _normalize_text(text)) if line.strip()]


def _label_variants(label: str) -> str:
    """生成匹配标签空格变体的正则，如 "甲方" -> "甲\\s*方" 同时匹配 "甲方" 和 "甲 方"."""
    if not label:
        return ""
    if re.search(r'[一-鿿㐀-䶿]', label):
        parts = [re.escape(ch) for ch in label]
        return r'\s*'.join(parts)
    return re.escape(label)


def _clean_ocr_text(text: str) -> str:
    """清理 LibreOffice 转换或 PaddleOCR 输出中的常见噪声."""
    text = _normalize_text(text)
    # 移除 CJK 字符之间的多余空格（WPS 格式常见问题）
    text = re.sub(r'(?<=[一-鿿])\s{1,3}(?=[一-鿿])', '', text)
    # 统一全角符号为半角
    text = text.replace('（', '(').replace('）', ')').replace('：', ':').replace('，', ',')
    # 归一化多余空行
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text


def _extract_raw_binary_text(path: Path) -> str:
    """从二进制 .doc 文件暴力提取可读文本（LibreOffice 转换失败时的兜底方案）."""
    try:
        raw = path.read_bytes()
    except OSError:
        return ""

    # 尝试多种编码解码
    for encoding in ('gbk', 'gb2312', 'gb18030', 'utf-16-le', 'utf-8'):
        try:
            text = raw.decode(encoding, errors='strict')
            lines = []
            for line in text.split('\n'):
                line = line.strip()
                if len(line) < 2:
                    continue
                cjk_ratio = sum(1 for c in line if '一' <= c <= '鿿') / max(len(line), 1)
                if cjk_ratio > 0.1 or len(line) > 20:
                    lines.append(line)
            if len(lines) > 5:
                return '\n'.join(lines)
        except (UnicodeDecodeError, LookupError):
            continue

    # 最后手段：按 UTF-16LE 逐码元提取中英文
    result_chars: list[str] = []
    for i in range(0, len(raw) - 1, 2):
        try:
            code = struct.unpack_from('<H', raw, i)[0]
        except struct.error:
            break
        if (0x4e00 <= code <= 0x9fff) or (0x3400 <= code <= 0x4dbf) or \
           (0x0020 <= code <= 0x007e) or code in (0x000a, 0x000d):
            result_chars.append(chr(code))
        else:
            if result_chars and result_chars[-1] != '\n':
                result_chars.append('\n')
    return ''.join(result_chars)


def _contains_any(text: str, labels: list[str]) -> bool:
    return any(label in text for label in labels)


def _payment_amount_candidates(text: str) -> list[str]:
    values = []
    for value in _money_values(text):
        if value not in values:
            values.append(value)
    return values


def _extract_payment_amount(text: str) -> tuple[str, list[str], str]:
    """Extract the actual payment amount from generic voucher text.

    The rule is label/table driven. It never relies on a fixed sample value, file
    name, or image position.
    """
    lines = _line_values(text)
    target_labels = ["本次回款金额", "本次收款金额", "本次付款金额", "本次到账金额", "回款金额", "收款金额"]
    invoice_amount_labels = ["发票金额", "票面金额", "价税合计"]
    all_candidates = _payment_amount_candidates(text)

    for index, line in enumerate(lines):
        if not _contains_any(line, target_labels):
            continue
        target_pos = min((line.find(label) for label in target_labels if label in line), default=-1)
        same_line = []
        for label in target_labels:
            if label in line:
                same_line = _money_values(line.split(label, 1)[1])
                break
        if same_line:
            return same_line[0], all_candidates, ""

        header_window = lines[max(0, index - 2): index + 1]
        prior_amount_columns = 0
        for item in header_window:
            if item == line and target_pos >= 0:
                prior_amount_columns += sum(1 for label in invoice_amount_labels if 0 <= item.find(label) < target_pos)
            elif _contains_any(item, invoice_amount_labels):
                prior_amount_columns += 1
        money_after: list[str] = []
        for next_line in lines[index + 1: index + 12]:
            money_after.extend(_money_values(next_line))
        if len(money_after) > prior_amount_columns:
            return money_after[prior_amount_columns], all_candidates, ""

    for label in target_labels:
        amount = _extract_amount_by_labels(text, [label], window=120, prefer="first")
        if amount:
            return amount, all_candidates, ""

    if len(all_candidates) == 1:
        return all_candidates[0], all_candidates, ""
    return "", all_candidates, "未能可靠定位本次回款金额，请人工确认"


def _section_between(text: str, start_patterns: list[str], end_patterns: list[str]) -> str:
    text = _normalize_text(text)
    start = None
    for pattern in start_patterns:
        match = re.search(pattern, text, re.I)
        if match and (start is None or match.start() < start.start()):
            start = match
    if not start:
        return ""
    end_index = len(text)
    for pattern in end_patterns:
        match = re.search(pattern, text[start.end():], re.I)
        if match:
            end_index = min(end_index, start.end() + match.start())
    return text[start.end():end_index]


def _extract_section_name(section: str) -> str:
    section = _normalize_text(section)
    variant = _label_variants("名称")
    for line in re.split(r"[\n\r]+", section):
        match = re.search(rf"{variant}\s*[:：]\s*(.+)", line.strip())
        if match:
            value = _clean_value(match.group(1))
            if value:
                return value
    match = re.search(rf"{variant}\s*[:：]\s*([^\n\r]+)", section)
    return _clean_value(match.group(1)) if match else ""


def _invoice_names(text: str) -> list[str]:
    names = []
    variant = _label_variants("名称")
    for match in re.finditer(rf"{variant}\s*[:：]\s*([^\n\r]+)", _normalize_text(text)):
        value = _clean_value(match.group(1))
        if value and value not in names:
            names.append(value)
    return names


def _extract_invoice_party(text: str, role: str) -> str:
    names = _invoice_names(text)
    if role == "buyer":
        section = _section_between(
            text,
            [r"购买方", r"购\s*买\s*方", r"买方"],
            [r"密码区", r"货物或应税劳务", r"货物或应税务", r"规格型号", r"销售方", r"销\s*售\s*方"],
        )
        fallback_labels = ["购买方名称", "购方名称", "买方名称", "购买方", "购方", "买方"]
        fallback_name = names[0] if names else ""
    else:
        section = _section_between(
            text,
            [r"销售方", r"销\s*售\s*方", r"销方", r"卖方"],
            [r"备注", r"收款人", r"复核", r"开票人", r"销售方\s*[:：]?\s*[（(]?章"],
        )
        fallback_labels = ["销售方名称", "销方名称", "卖方名称", "销售方", "销方", "卖方"]
        fallback_name = names[1] if len(names) > 1 else ""
    return _extract_section_name(section) or fallback_name or _match_after(text, fallback_labels)


def _parse_chinese_section(section: str) -> int:
    total = 0
    number = 0
    for char in section:
        if char in _CN_DIGITS:
            number = _CN_DIGITS[char]
        elif char in _CN_UNITS:
            total += (number or 1) * _CN_UNITS[char]
            number = 0
    return total + number


def _parse_chinese_integer(value: str) -> int:
    value = re.sub(r"[人民币圆元整正\s]", "", value)
    value = re.split(r"[角分]", value, 1)[0]
    total = 0
    if "亿" in value:
        before, value = value.split("亿", 1)
        total += _parse_chinese_section(before) * 100000000
    if "万" in value:
        before, value = value.split("万", 1)
        total += _parse_chinese_section(before) * 10000
    total += _parse_chinese_section(value)
    return total


def _extract_chinese_amount(text: str) -> str:
    amount_chars = "零〇一壹二贰两三叁四肆五伍六陆七柒八捌九玖十拾百佰千仟万亿圆元角分整正"
    patterns = [
        rf"(?:合同金额|合同总金额|总金额|价款|金额|人民币)[:：\s]*(?:人民币)?\s*([{amount_chars}]{{2,}})",
        rf"人民币\s*([{amount_chars}]{{2,}})",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if not match:
            continue
        amount = _parse_chinese_integer(match.group(1))
        if amount:
            return f"{Decimal(amount).quantize(Decimal('0.01'))}"
    return ""


def _extract_amount(text: str, labels: list[str] | None = None) -> str:
    text = _normalize_text(text)
    labels = labels or ["合同金额", "合同总金额", "总金额", "价款", "金额", "人民币"]
    label_expr = "|".join(re.escape(label) for label in labels)
    labeled_patterns = [
        rf"(?:{label_expr})[:：\s]*(?:人民币)?[^\n\r\d¥￥]{{0,40}}[（(]?\s*[¥￥]\s*([0-9][0-9,]*(?:\.\d{{1,2}})?)",
        rf"(?:{label_expr})[:：\s]*(?:人民币)?\s*([0-9][0-9,]*(?:\.\d{{1,2}})?)",
        rf"(?:{label_expr})[:：\s]*(?:人民币)?[^\n\r\d]{{0,80}}([0-9][0-9,]*(?:\.\d{{1,2}})?)",
    ]
    for pattern in labeled_patterns:
        match = re.search(pattern, text, re.I)
        if match:
            amount = _decimal_text(match.group(1))
            if amount:
                return amount

    symbol_matches = re.findall(r"[¥￥]\s*([0-9][0-9,]*(?:\.\d{1,2})?)", text)
    if symbol_matches:
        amounts = [_decimal_text(item) for item in symbol_matches]
        amounts = [item for item in amounts if item]
        if amounts:
            return max(amounts, key=lambda item: Decimal(item))

    chinese = _extract_chinese_amount(text)
    if chinese:
        return chinese

    candidates = []
    for match in re.finditer(r"\b([0-9][0-9,]*(?:\.\d{1,2})?)\b", text):
        raw = match.group(1)
        if re.fullmatch(r"20\d{2}", raw) or ("." not in raw and "," not in raw):
            continue
        amount = _decimal_text(raw)
        if amount:
            candidates.append(amount)
    return max(candidates, key=lambda item: Decimal(item)) if candidates else ""


def _extract_contract_total_amount(text: str) -> str:
    text = _normalize_text(text)
    priority_patterns = [
        # "本合同总金额：人民币小写：1000000.00元" / "本合同总金额：人民币（小写）：1000000.00元"
        r"(?:本合同总金额|合同总金额|合同金额|合同价款|总价款|总金额)\s*[:：]\s*(?:人民币)?\s*(?:[（(]?小写[）)]?)?\s*[:：]?\s*(?:人民币)?\s*[¥￥]?\s*([0-9][0-9,]*(?:\.\d{1,2})?)",
        # 字段名后面直接跟数字
        r"(?:本合同总金额|合同总金额|合同金额|合同价款|总价款|总金额)\s*[:：]?\s*(?:人民币)?\s*(?:小写)?\s*[:：]?\s*[¥￥]?\s*([0-9][0-9,]*(?:\.\d{1,2})?)",
        # 字段名后面不远处有小写+¥符号
        r"(?:本合同总金额|合同总金额|合同金额|合同价款|总价款|总金额)[^\n\r]{0,80}?(?:小写)[^\n\r]{0,20}?[¥￥]?\s*([0-9][0-9,]*(?:\.\d{1,2})?)",
        # 字段名后面不远处有¥符号
        r"(?:本合同总金额|合同总金额|合同金额|合同价款|总价款|总金额)[^\n\r]{0,60}?[¥￥]\s*([0-9][0-9,]*(?:\.\d{1,2})?)",
    ]
    for pattern in priority_patterns:
        match = re.search(pattern, text, re.I)
        if match:
            amount = _decimal_text(match.group(1))
            if amount:
                return amount
    heading_match = re.search(r"(?:合同金额及付款方式|合同金额|合同价款)([\s\S]{0,260})", text)
    if heading_match:
        snippet = heading_match.group(1)
        for label in ["本合同总金额", "合同总金额", "合同金额", "人民币小写", "人民币（小写）", "小写"]:
            amount = _extract_amount_by_labels(snippet, [label], window=120, prefer="first")
            if amount:
                return amount
    return ""


def _has_explicit_contract_total_label(text: str) -> bool:
    text = _normalize_text(text)
    labels = ["本合同总金额", "合同总金额", "合同金额", "合同价款", "总价款", "总金额"]
    pattern = "|".join(_label_variants(l) for l in labels)
    return bool(re.search(f"(?:{pattern})", text))


def _extract_date(text: str) -> str:
    text = _normalize_text(text)
    match = re.search(r"(20\d{2})\s*[年\-/\.]\s*(\d{1,2})\s*[月\-/\.]\s*(\d{1,2})\s*日?", text)
    if not match:
        return ""
    return f"{match.group(1)}-{int(match.group(2)):02d}-{int(match.group(3)):02d}"


def _match_after(text: str, labels: list[str], limit: int = 120) -> str:
    text = _normalize_text(text)
    lines = [line.strip() for line in re.split(r"[\n\r]+", text)]
    for index, line in enumerate(lines):
        normalized = line.strip()
        for label in labels:
            variant = _label_variants(label)
            match = re.search(rf"{variant}\s*(?:[（(][^）)]*[）)])?\s*[:：]?\s*(.*)$", normalized)
            if match:
                value = _clean_value(match.group(1), limit)
                if value and value not in {"甲方", "乙方", "章"}:
                    return value
                for next_line in lines[index + 1:]:
                    next_value = _clean_value(next_line, limit)
                    if next_value:
                        return next_value
    for label in labels:
        variant = _label_variants(label)
        match = re.search(rf"{variant}\s*(?:[（(][^）)]*[）)])?\s*[:：]?\s*([^\n\r;；]{{1,{limit}}})", text)
        if match:
            value = _clean_value(match.group(1), limit)
            if value not in {"甲方", "乙方", "章"}:
                return value
    return ""


def _extract_contract_no(text: str) -> str:
    text = _normalize_text(text)
    # 优先匹配 HT-xxxx 格式的标准合同编号
    match = re.search(r"\b(HT[-_A-Za-z0-9]{4,})\b", text, re.I)
    if match:
        return match.group(1)
    # 优先取明确的合同编号标签
    labeled = _match_after(text, ["合同编号", "合同号"], 80)
    if labeled:
        match = re.search(r"[A-Za-z0-9][A-Za-z0-9\-_]{3,}", labeled)
        return match.group(0) if match else labeled
    # 兜底：无合同编号时使用项目编号作为标识
    fallback = _match_after(text, ["项目编号", "项目号", "编号"], 80)
    if fallback:
        match = re.search(r"[A-Za-z0-9][A-Za-z0-9\-_]{3,}", fallback)
        return match.group(0) if match else fallback
    return ""


def _infer_project_name(text: str) -> str:
    lines = _line_values(text)
    candidates: list[str] = []
    for index, line in enumerate(lines[:40]):
        cleaned = _clean_value(line, 160)
        if not cleaned:
            continue
        if "合同" in cleaned and any(word in cleaned for word in ["项目", "系统", "服务", "运维", "建设", "开发"]):
            candidates.append(cleaned)
        if "订立本合同" in cleaned:
            match = re.search(r"就(.+?)(?:订立|签订|达成)本合同", cleaned)
            if match:
                candidates.append(_clean_value(match.group(1), 160))
        if index + 1 < len(lines) and any(word in cleaned for word in ["年度", "项目", "系统"]) and "合同" not in cleaned:
            next_line = _clean_value(lines[index + 1], 120)
            if "合同" in next_line:
                candidates.append(_clean_value(f"{cleaned}{next_line}", 160))
    for candidate in candidates:
        value = re.sub(r"^(?:预算管理一体化系统)?", "", candidate).strip()
        value = re.sub(r"(?:合同|协议|书)$", "", value).strip()
        if len(value) >= 6 and value not in {"合同", "项目合同"}:
            return value
    return ""


def _extract_invoice_no(text: str) -> str:
    text = _normalize_text(text)
    labeled = _match_after(text, ["发票号码", "发票号", "号码"], 80)
    if labeled:
        match = re.search(r"[A-Za-z0-9-]{6,30}", labeled)
        return match.group(0) if match else labeled
    match = re.search(r"(?:发票号码|发票号|No\.?)\s*[:：]?\s*([A-Za-z0-9-]{6,30})", text, re.I)
    return match.group(1) if match else ""


def _extract_contract(text: str) -> dict:
    text = _normalize_text(text)
    party_a = _match_after(text, ["甲方", "委托方", "采购方", "客户名称", "客户"])
    party_b = _match_after(text, ["乙方", "承包方", "服务方", "供应商"])
    party_c = _match_after(text, ["丙方", "监理方", "第三方"])
    return {
        "project_name": _match_after(text, ["项目名称", "项目名"]) or _infer_project_name(text),
        "contract_amount": _extract_contract_total_amount(text) or _extract_amount(text, ["本合同总金额", "合同总金额", "合同金额", "总金额", "价款", "金额", "人民币"]),
        "contract_no": _extract_contract_no(text),
        "sign_date": _extract_date(text),
        "party_a": party_a,
        "party_b": party_b,
        "party_c": party_c,
        "customer": party_a,
        "project_type": _match_after(text, ["项目类型", "业务类型", "服务类型", "合同类型"], 50),
    }


def _contract_manual_required(info: dict) -> list[str]:
    labels = {
        "project_name": "项目名称",
        "contract_amount": "合同金额",
        "party_a": "甲方/客户",
        "party_b": "乙方",
    }
    return [label for key, label in labels.items() if not info.get(key)]


def _call_llm_contract_extract(text: str, config: LLMRuntimeConfig) -> dict:
    if not config.enabled or not config.api_key:
        return {}
    base_url = config.api_base_url.rstrip("/")
    url = f"{base_url}/chat/completions"
    prompt = (
        '你是合同信息抽取助手。请只返回 JSON，不要 Markdown。\n\n'
        '从以下合同文本中提取字段，所有字段无法确定时返回空字符串：\n'
        '- project_name: 项目名称（优先取封面标题，去掉末尾"合同"/"协议"字样。也可从"就……项目订立本合同"中提取。）\n'
        '- contract_amount: 合同金额，只取"本合同总金额"或"合同总金额"对应的数字，单位为元，不要取分期付款、保证金、罚款、税款、账号、日期或编号。\n'
        '- contract_no: 合同编号（优先"合同编号"/"合同号"；若无，可使用"项目编号"/"项目号"作为兜底；都不要时才返回空字符串）。\n'
        '- sign_date: 签订日期，格式 YYYY-MM-DD。\n'
        '- party_a: 甲方（委托方/采购方）的单位名称或个人姓名，只从"甲方"/"委托方"/"采购方"标签后提取。\n'
        '- party_b: 乙方（承包方/服务方）的单位名称或个人姓名，只从"乙方"/"承包方"/"服务方"标签后提取。\n'
        '- party_c: 丙方（监理方/第三方），只从"丙方"/"监理方"/"第三方"标签后提取。\n'
        '- project_type: 项目类型。\n\n'
        '重要规则：\n'
        '1. 标签与值之间可能有多余空格（如"甲 方"），请宽容匹配。\n'
        '2. 金额格式多样：可能是"1000000.00元"、"¥1,000,000.00"、"壹佰万元整"。\n'
        '3. 甲方/乙方可能包含单位名称，如文本中出现"甲方：某某公司（盖章）"，提取"某某公司"。\n'
        '4. 如有多个金额，优先取标有"总金额"/"合同总金额"的那个。\n'
        '5. 文本可能来自扫描件 OCR，含少量错字（如"己方"→"乙方"、"合司"→"合同"），请容错识别。\n\n'
        '示例 - WPS 格式合同：\n'
        '输入: "2024年度xx财政预算管理一体化系统运行维护项目（二次）合同\\n甲 方：新疆杜云飞\\n乙 方：栗姜涛\\n本合同总金额：人民币小写：1000000.00元。"\n'
        '输出: {"project_name":"2024年度xx财政预算管理一体化系统运行维护项目（二次）","contract_amount":"1000000.00","contract_no":"","sign_date":"","party_a":"新疆杜云飞","party_b":"栗姜涛","party_c":"","project_type":""}'
    )
    payload = {
        "model": config.model,
        "messages": [
            {"role": "system", "content": prompt},
            {"role": "user", "content": _normalize_text(text)[:15000]},
        ],
        "temperature": 0,
        "response_format": {"type": "json_object"},
    }
    headers = {"Authorization": f"Bearer {config.api_key}", "Content-Type": "application/json"}
    response = httpx.post(url, headers=headers, json=payload, timeout=45)
    response.raise_for_status()
    content = response.json()["choices"][0]["message"]["content"]
    data = json.loads(content)
    cleaned = {
        "project_name": _clean_value(data.get("project_name"), 160),
        "contract_amount": _decimal_text(str(data.get("contract_amount") or "")),
        "contract_no": _clean_value(data.get("contract_no"), 80),
        "sign_date": _extract_date(str(data.get("sign_date") or "")) or _clean_value(data.get("sign_date"), 20),
        "party_a": _clean_value(data.get("party_a"), 120),
        "party_b": _clean_value(data.get("party_b"), 120),
        "party_c": _clean_value(data.get("party_c"), 120),
        "project_type": _clean_value(data.get("project_type"), 50),
    }
    cleaned["customer"] = cleaned["party_a"]
    return cleaned

def _complete_contract_info(db: Session, raw_text: str, info: dict) -> tuple[dict, bool, dict, list[str], str]:
    merged = dict(info)
    field_sources = {key: "rule" for key, value in merged.items() if value}
    llm_used = False
    llm_error = ""
    config = get_llm_runtime_config(db)
    has_total_label = _has_explicit_contract_total_label(raw_text)
    if config.enabled and config.api_key:
        try:
            llm_info = _call_llm_contract_extract(raw_text, config)
            if llm_info:
                llm_used = True
                for key, value in llm_info.items():
                    if not value:
                        continue
                    if key == "contract_no" and merged.get(key):
                        continue
                    if key == "contract_amount" and merged.get(key) and has_total_label:
                        continue
                    if key in {"sign_date", "party_a", "party_b", "party_c", "customer"} and merged.get(key):
                        continue
                    if key == "project_name" and merged.get(key) and len(str(merged[key])) >= len(str(value)):
                        continue
                    if key in {"project_name", "contract_amount", "sign_date", "party_a", "party_b", "party_c", "project_type", "customer"}:
                        merged[key] = value
                        field_sources[key] = "llm"
                    elif not merged.get(key):
                        merged[key] = value
                        field_sources[key] = "llm"
        except Exception as exc:
            llm_error = str(exc)
    manual_required = _contract_manual_required(merged)
    return merged, llm_used, field_sources, manual_required, llm_error


def _extract_invoice(text: str) -> dict:
    text = _normalize_text(text)
    # 发票票面中“金额”列、合同备注和价税合计会同时出现，必须按票据字段优先级取值。
    amount = _extract_amount_by_labels(
        text,
        ["价税合计（小写）", "价税合计(小写)", "价税合计", "发票金额（含税）", "发票金额(含税)", "发票金额"],
        prefer="max",
    )
    tax_amount = _extract_amount_by_labels(text, ["合计税额", "税额"], prefer="first")
    amount_without_tax = _extract_amount_by_labels(
        text,
        ["不含税金额", "合计金额（不含税）", "合计金额(不含税)", "金额"],
        prefer="first",
    )

    all_money = _money_values(text)
    if not amount and all_money:
        amount = max(all_money, key=Decimal)

    tax_rate = ""
    rate_match = re.search(r"(\d+(?:\.\d+)?)\s*%", text)
    if rate_match:
        tax_rate = rate_match.group(1)
        after_rate_values = _money_values(text[rate_match.end(): rate_match.end() + 100])
        if after_rate_values:
            tax_amount = after_rate_values[0]

    try:
        amount_decimal = _to_decimal(amount)
        tax_decimal = _to_decimal(tax_amount)
        without_tax_decimal = _to_decimal(amount_without_tax)

        if not tax_decimal and tax_rate and without_tax_decimal:
            tax_decimal = (without_tax_decimal * Decimal(tax_rate) / Decimal("100")).quantize(Decimal("0.01"))
            tax_amount = f"{tax_decimal}"
        if not without_tax_decimal and amount_decimal and tax_decimal:
            without_tax_decimal = (amount_decimal - tax_decimal).quantize(Decimal("0.01"))
            amount_without_tax = f"{without_tax_decimal}"
        if not without_tax_decimal and amount_decimal and tax_rate:
            rate = Decimal(tax_rate) / Decimal("100")
            without_tax_decimal = (amount_decimal / (Decimal("1") + rate)).quantize(Decimal("0.01"))
            amount_without_tax = f"{without_tax_decimal}"
            tax_decimal = (amount_decimal - without_tax_decimal).quantize(Decimal("0.01"))
            tax_amount = f"{tax_decimal}"
        if not amount_decimal and without_tax_decimal and tax_decimal:
            amount_decimal = (without_tax_decimal + tax_decimal).quantize(Decimal("0.01"))
            amount = f"{amount_decimal}"
        if not amount and amount_without_tax and tax_amount:
            amount = f"{(Decimal(amount_without_tax) + Decimal(tax_amount)).quantize(Decimal('0.01'))}"
        if not tax_amount and amount and amount_without_tax:
            tax_amount = f"{(Decimal(amount) - Decimal(amount_without_tax)).quantize(Decimal('0.01'))}"
        if not tax_rate and tax_amount and amount_without_tax and Decimal(amount_without_tax) > 0:
            tax_rate = f"{(Decimal(tax_amount) / Decimal(amount_without_tax) * Decimal('100')).quantize(Decimal('0.01'))}"
    except InvalidOperation:
        pass

    return {
        "invoice_no": _extract_invoice_no(text),
        "invoice_code": _match_after(text, ["发票代码"], 80),
        "amount": amount,
        "amount_without_tax": amount_without_tax,
        "tax_rate": tax_rate,
        "tax_amount": tax_amount,
        "invoice_date": _extract_date(text),
        "buyer": _extract_invoice_party(text, "buyer"),
        "seller": _extract_invoice_party(text, "seller"),
        "project_name": _match_after(text, ["项目名称"], 120),
        "contract_no": _extract_contract_no(text),
    }


def _extract_payment(text: str) -> dict:
    text = _normalize_text(text)
    payer = _match_after(text, ["付款方", "付款人", "付款账户名", "付款单位", "汇款方"])
    payee = _match_after(text, ["收款方", "收款人", "收款账户名", "收款单位"])
    serial = _match_after(text, ["凭证编号", "流水号", "交易流水号", "回单编号", "凭证号"], 80)
    if not serial:
        serial_match = re.search(r"\b[A-Z]{1,8}[-_]\d{4}[-_A-Za-z0-9]{3,}\b", text)
        serial = serial_match.group(0) if serial_match else ""
    amount, amount_candidates, amount_warning = _extract_payment_amount(text)
    invoice_no = _extract_invoice_no(text)
    remark_parts = [item for item in [serial, payer, payee] if item]
    return {
        "amount": amount,
        "amount_candidates": amount_candidates,
        "amount_warning": amount_warning,
        "payment_date": _extract_date(text),
        "payer": payer,
        "payee": payee,
        "serial_no": serial,
        "invoice_no": invoice_no,
        "remark": " / ".join(remark_parts),
    }


def _confidence_for(info: dict, base: float) -> dict:
    return {key: (base if value else 0.0) for key, value in info.items()}


def _parse_docx(path: Path) -> str:
    document = Document(str(path))
    lines = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                lines.append(" ".join(values))
    return "\n".join(lines)


def _convert_doc_with_libreoffice(path: Path) -> Path:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError("当前环境未安装 LibreOffice，无法解析 .doc 文件，请转换为 .docx 后上传")
    temp_dir = Path(tempfile.mkdtemp(prefix="doc_convert_"))
    env = os.environ.copy()
    if "HOME" not in env or not env["HOME"]:
        env["HOME"] = str(temp_dir)

    def _try_convert(args: list[str]) -> subprocess.CompletedProcess:
        return subprocess.run(args, capture_output=True, text=True, timeout=120, env=env)

    # 尝试 1：标准转换
    cmd = [soffice, "--headless", "--convert-to", "docx", "--outdir", str(temp_dir), str(path)]
    proc = _try_convert(cmd)

    # 尝试 2：MS Word 97 输入过滤器（WPS 格式兼容）
    if proc.returncode != 0:
        filter_cmd = [soffice, "--headless", "--infilter=MS Word 97", "--convert-to", "docx", "--outdir", str(temp_dir), str(path)]
        proc = _try_convert(filter_cmd)

    if proc.returncode != 0:
        shutil.rmtree(temp_dir, ignore_errors=True)
        raise RuntimeError(
            proc.stderr.strip() or proc.stdout.strip() or
            ".doc 转换失败，请将文件另存为 .docx 格式后重新上传"
        )

    converted = temp_dir / f"{path.stem}.docx"
    if not converted.exists():
        matches = list(temp_dir.glob("*.docx"))
        if not matches:
            shutil.rmtree(temp_dir, ignore_errors=True)
            raise RuntimeError(".doc 转换后未生成 docx 文件，请将文件另存为 .docx 格式后重新上传")
        converted = matches[0]
    return converted


def _read_document_text(path: Path) -> str:
    if path.suffix.lower() == ".doc":
        try:
            converted = _convert_doc_with_libreoffice(path)
        except RuntimeError:
            # LibreOffice 转换失败 → 尝试二进制暴力提取
            raw_text = _extract_raw_binary_text(path)
            if raw_text:
                return raw_text
            raise
        try:
            return _parse_docx(converted)
        finally:
            shutil.rmtree(converted.parent, ignore_errors=True)
    if path.suffix.lower() == ".docx":
        return _parse_docx(path)
    return path.read_text(encoding="utf-8", errors="ignore") if path.exists() else path.name


def _call_paddleocr(path: Path, recognition_type: str) -> tuple[str, float, list]:
    settings = get_settings()
    with path.open("rb") as file:
        files = {"file": (path.name, file, "application/octet-stream")}
        response = httpx.post(settings.paddleocr_url, params={"type": recognition_type}, files=files, timeout=90)
    response.raise_for_status()
    payload = response.json()
    data = payload.get("data", payload)
    if payload.get("code") not in {None, 200, "200"}:
        raise RuntimeError(data.get("error") or payload.get("message") or "PaddleOCR 识别失败")

    raw_lines = data.get("lines") or []
    raw_text_field = data.get("raw_text") or data.get("text") or ""
    texts: list[str] = []
    confidences: list[float] = []
    for line in raw_lines:
        if isinstance(line, dict):
            if not raw_text_field:
                texts.append(str(line.get("text") or ""))
            if line.get("confidence") is not None:
                confidences.append(float(line["confidence"]))
        elif isinstance(line, (list, tuple)) and line:
            if not raw_text_field:
                texts.append(str(line[0]))
            if len(line) >= 3:
                try:
                    confidences.append(float(line[2]))
                except (TypeError, ValueError):
                    pass
        elif not raw_text_field:
            texts.append(str(line))
    raw_text = raw_text_field if isinstance(raw_text_field, str) and raw_text_field else "\n".join(t for t in texts if t)
    confidence = sum(confidences) / len(confidences) if confidences else float(data.get("confidence") or 0)
    return raw_text, confidence, raw_lines


def _extract_by_type(raw_text: str, recognition_type: str) -> dict:
    if recognition_type == "invoice":
        return _extract_invoice(raw_text)
    if recognition_type == "payment":
        return _extract_payment(raw_text)
    return _extract_contract(raw_text)


def recognize_file(db: Session, file_path: str, recognition_type: str) -> dict:
    start = time.perf_counter()
    path = Path(file_path)
    suffix = path.suffix.lower()
    engine = "parser" if suffix in {".doc", ".docx"} else "paddleocr"
    try:
        if engine == "parser":
            raw_text = _read_document_text(path)
            raw_result = {"text": raw_text[:2000]}
            base_confidence = 0.95 if raw_text else 0.0
        else:
            raw_text, base_confidence, lines = _call_paddleocr(path, recognition_type)
            raw_result = {"lines": lines, "text": raw_text[:2000]}

        info = _extract_by_type(raw_text, recognition_type)
        llm_used = False
        field_sources = {key: "rule" for key, value in info.items() if value}
        manual_required_fields: list[str] = []
        llm_error = ""
        if recognition_type == "contract":
            info, llm_used, field_sources, manual_required_fields, llm_error = _complete_contract_info(db, raw_text, info)
        confidence = _confidence_for(info, base_confidence)
        duration = time.perf_counter() - start
        status = "success" if raw_text else "manual_required"
        if recognition_type == "contract" and raw_text and manual_required_fields:
            status = "manual_required"
        if recognition_type == "payment" and raw_text and not info.get("amount"):
            status = "manual_required"
        log = OcrRecognitionLog(
            file_path=file_path,
            file_name=path.name,
            recognition_type=recognition_type,
            engine=engine,
            raw_result=json.dumps(raw_result, ensure_ascii=False),
            extracted_info=json.dumps(info, ensure_ascii=False),
            confidence=base_confidence,
            status=status,
            duration=duration,
        )
        db.add(log)
        db.flush()
        result = {
            "log_id": log.id,
            "raw_text": raw_text[:2000],
            "extracted_info": info,
            "confidence": confidence,
            "engine": engine,
            "duration": duration,
            "status": status,
            "parse_source": "llm" if llm_used else "rule",
            "llm_used": llm_used,
            "field_sources": field_sources,
            "manual_required_fields": manual_required_fields,
        }
        if llm_error:
            result["llm_error"] = llm_error
        return result
    except Exception as exc:
        duration = time.perf_counter() - start
        log = OcrRecognitionLog(
            file_path=file_path,
            file_name=path.name,
            recognition_type=recognition_type,
            engine=engine,
            extracted_info=json.dumps({}, ensure_ascii=False),
            confidence=0,
            status="failed",
            duration=duration,
            error_message=str(exc),
        )
        db.add(log)
        db.flush()
        return {"log_id": log.id, "raw_text": "", "extracted_info": {}, "confidence": {}, "engine": engine, "duration": duration, "status": "failed", "error_message": str(exc)}
